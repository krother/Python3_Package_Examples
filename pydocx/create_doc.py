
# Create a Word document:

from docx import Document

d = Document()

d.add_heading('Hamlet')
d.add_heading('dramatis personae', 2)
d.add_paragraph('Hamlet, the Prince of Denmark')

d.save('hamlet.docx')
